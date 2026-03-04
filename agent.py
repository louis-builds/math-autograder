from typing import List
from google.genai import types
from google.adk.agents import Agent
from google.adk.runners import Runner
from google.adk.sessions import InMemorySessionService

# 1. Define an exact calculator tool (Core: Prevent AI calculation errors)
def strict_calculator(expression: str) -> str:
    """Used to execute addition, subtraction, multiplication, and division calculations for New Zealand exam papers.
    Supported operators include: +, -, *, /, x, ×, ÷
    """
    # Clean up and replace common multiplication and division expressions
    exp = expression.replace('x', '*').replace('×', '*').replace('÷', '/')
    try:
        # In an actual production environment, it's recommended to use safe math or ast.literal_eval combinations, or the sympy library instead of eval
        result = float(eval(exp))
        # Remove trailing .0 from integer results
        if result.is_integer():
            return str(int(result))
        return str(result)
    except Exception as e:
        return f"Calculation Error: {e}"

# 2. Define an Agent specifically for parsing PDFs
pdf_parser = Agent(
    name="NZ_Paper_Scanner",
    model="gemini-2.5-flash",
    instruction="""You are responsible for reading the PDF/image content of New Zealand primary and secondary school math exam papers.
Responsibilities:
1. Accurately identify all addition, subtraction, multiplication, and division expressions, as well as fractions and percentages within them.
2. You must retain the original question numbers (e.g., Q1, Q2, etc.).
3. If you encounter columnar arithmetic problems on the page, you need to convert them into standard horizontal expressions.
4. Output the organized list of standard expressions completely and clearly to the Math_Expert in the team."""
)

# 3. Define the Solver Agent
math_solver = Agent(
    name="Math_Expert",
    model="gemini-2.5-flash",
    tools=[strict_calculator],
    instruction="""You are a Math Solver Expert. You receive organized math expressions from the scanner.
Responsibilities:
1. For every expression received, you are strictly forbidden to "mentally calculate" or guess the answer yourself.
2. You must call the `strict_calculator` tool, input the mathematical expression to calculate it, and obtain a 100% accurate forward calculation result.
3. Consolidate the original question number, original expression, and the forward result calculated by the tool, and hand it over to the Math_Checker for verification."""
)

# 4. Define the Checker Agent
math_checker = Agent(
    name="Math_Checker",
    model="gemini-2.5-flash",
    tools=[strict_calculator],
    instruction="""You are a rigorous Checker. Your core responsibility is to perform strict reverse verification on the math answers provided by the Solver to completely eliminate errors.
The verification logic is as follows (you must use reverse operations to verify):
- If the original expression is subtraction (e.g., 100 - 34 = 66), you must verify the corresponding addition (call the tool to calculate 66 + 34 and see if it equals 100).
- If the original expression is addition (e.g., 66 + 34 = 100), you must verify the corresponding subtraction (call the tool to calculate 100 - 34 and see if it equals 66).
- If the original expression is division (e.g., 100 / 4 = 25), you must verify the corresponding multiplication (call the tool to calculate 25 * 4 and see if it equals 100).
- If the original expression is multiplication (e.g., 25 * 4 = 100), you must verify the corresponding division (call the tool to calculate 100 / 4 and see if it equals 25).

Workflow:
1. Analyze the expression and initial result sent by the Solver.
2. Construct the corresponding "reverse expression".
3. Call `strict_calculator` to calculate the result of the reverse expression.
4. Compare and verify:
   - If the reverse calculation result matches the original value on the other side of the equation, it means the initial calculation was correct. Output: "Verification Passed ✅" and provide the accurate answer in the final report.
   - If it does not match, it indicates the initial calculation or parsing was wrong. Output: "Verification Failed ❌", and correct the answer."""
)

# New: Local PDF Reading Tool
def extract_text_from_pdf(pdf_path: str) -> str:
    """If the user provides a local PDF file path, call this tool to extract text content.
    Args:
        pdf_path (str): The absolute or relative path to the local PDF file on the computer.
    Returns:
        str: The extracted text exam paper content. If it fails, return an error message.
    """
    try:
        import PyPDF2
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
    except ImportError:
        return "Cannot read: Please remind the user to run `pip install PyPDF2` in the terminal first."
    except Exception as e:
        return f"Reading Error: File not found or corrupted ({e})"

# New: Export Results to Local Word Document Tool
def export_answers_to_doc(answers_text: str, output_path: str = "answers.docx") -> str:
    """If the user requests to generate a .doc or .docx file containing the answers, call this tool to export the final clean answers.
    Args:
        answers_text (str): Only contains the string content of the final answers, removing the verification thought process, keeping only question numbers and answers. Should match the original paper's layout format as much as possible (e.g., if questions are on the same line, answers should be on the same line, separated by spaces).
        output_path (str): The output file path, default is answers.docx.
    Returns:
        str: Information on whether export was successful or failed.
    """
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Math Assessment Answers', 0)
        
        # Write the incoming text line by line, letting the LLM control the layout style and spacing
        for line in answers_text.split('\n'):
            doc.add_paragraph(line)
                
        doc.save(output_path)
        return f"✅ Answers successfully exported to {output_path} matching your requested layout."
    except ImportError:
        return "Cannot export: Please remind the user to run `pip install python-docx` in the terminal first."
    except Exception as e:
        return f"Export failed: {e}"

# 5. Build the Supervisor Agent for the Team Agent
# (Since the current ADK cannot directly run complex teams via CLI, we can use a supervisor agent and provide it with other sub-agents as tools to orchestrate the process. However, for a simple CLI test, we create a composite agent here.)
root_agent = Agent(
    name="Math_Assessment_Master",
    model="gemini-2.5-flash",
    instruction="""You are the master of a math grading pipeline.
You need to sequentially call upon your knowledge to:
1. Obtain the math exam text or content provided by the user. If the user provides a path pointing to a file (e.g., ending in .pdf), first read it using `extract_text_from_pdf`.
2. Call tools to sequentially hand over processing to the Scanner -> Expert -> Checker.
However, in the current streamlined mode, since there's no native Team object support in the CLI, you need to concurrently handle these roles yourself.

Please strictly process the user's request according to the following steps:
Step 1 (Scanner): Identify addition, subtraction, multiplication, and division in the original problems and convert columnar formats to horizontal ones.
Step 2 (Expert): Use the strict_calculator tool to calculate the results from Step 1.
Step 3 (Checker): Use the strict_calculator based on inverse calculation rules to deduce backwards whether the answer from Step 2 is correct.
Step 4: Output the results that pass the final verification. If an error occurs, you must state the reason why the verification failed.
Step 5: If the user explicitly asks to generate a .doc or export final answers to a file, you must use the `export_answers_to_doc` tool to save the pure answer string (removing the "thought/verification process", e.g., keeping only the format "1. 1545") to a local file.
""",
    tools=[strict_calculator, extract_text_from_pdf, export_answers_to_doc]
)