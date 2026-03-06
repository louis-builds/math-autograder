import sys
import os
import io
import time # 必须引入 time 库
import fitz  # PyMuPDF
from typing import List
from pydantic import BaseModel, Field
from google import genai
from PIL import Image

from docx import Document
import dotenv

# 加载 .env 里的 API Key
dotenv.load_dotenv()

# --- 1. Data Structure Definitions ---
class ProblemEntry(BaseModel):
    page_number: int = Field(description="Page number")
    question_number: str = Field(description="The question label or number (e.g., '1)', '2.', '3'). If none, leave empty.", default="")
    question_type: str = Field(description="Type of question: 'calculation', 'multiple_choice', 'true_false', or 'other'")
    expression: str = Field(description="If question_type is 'calculation', put the left-hand math formula to compute (e.g., '(1/2)+0.5'). For other types, leave empty.", default="")
    llm_answer: str = Field(description="If question_type is NOT 'calculation', the LLM should directly provide the concise answer here (e.g., 'A', 'True', '5 m'). For calculations, leave empty.", default="")
    is_word_problem: bool = Field(description="Whether this is a word problem that requires logical reading before calculating")

class PaperStructure(BaseModel):
    problems: List[ProblemEntry]

# --- 2. Visual Conversion Logic (No poppler needed) ---
def convert_pdf_to_images(pdf_path: str):
    print(f"🚀 [1/4] Parsing PDF: {pdf_path}")
    try:
        doc = fitz.open(pdf_path)
        images = []
        for page in doc:
            # 2.0x scale ensures ~300DPI clarity for recognizing small symbols
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_data = pix.tobytes("png")
            images.append(Image.open(io.BytesIO(img_data)))
        return images
    except Exception as e:
        print(f"❌ Failed to read PDF: {e}")
        sys.exit(1)

# --- 3. Gemini Vision Parsing Core ---
def parse_with_gemini_vision(images: list) -> List[ProblemEntry]:
    print(f"👁️ [2/4] Using Gemini 2.5 Flash to recognize images ({len(images)} pages)...")
    
    api_key = os.environ.get("GEMINI_API_KEY")
    client = genai.Client(api_key=api_key)
    all_extracted_data = []

    for i, img in enumerate(images):
        page_num = i + 1
        print(f"  Recognizing page {page_num}...")
        
        prompt = """You are a math and test expert. Exhaustively extract ALL questions from the image.
        1. Extract the `question_number` if it exists as a prefix (e.g., '1)', '2.').
        2. Classify each question using `question_type` ('calculation', 'multiple_choice', 'true_false', 'other').
        3. If 'calculation': 
           - Convert columnar/text formats to standard math operators (+ - * /).
           - Output ONLY the left-hand side formula to `expression` (NO equals sign).
           - Convert fractions like '1 1/2' to '(1 + 1/2)'.
        4. If 'multiple_choice', 'true_false', or 'other':
           - DO NOT put anything in `expression`.
           - Instead, directly answer the question yourself and put your final concise answer (e.g., 'A', 'True', '300 grams') into the `llm_answer` field.
        5. Set `is_word_problem` to true for word problems, but false for plain variable substitutions like 'x+y'."""

        # --- Core: Use confirmed working model name ---
        # Add retry and slowdown logic
        for attempt in range(3):
            try:
                response = client.models.generate_content(
                    model='gemini-2.5-flash', 
                    contents=[prompt, img],
                    config={
                        'response_mime_type': 'application/json', 
                        'response_schema': PaperStructure
                    }
                )
                page_data = PaperStructure.model_validate_json(response.text)
                for prob in page_data.problems:
                    prob.page_number = page_num
                    all_extracted_data.append(prob)
                
                # Wait 5 seconds between pages to avoid 429 rate limits
                time.sleep(5) 
                break 

            except Exception as e:
                error_str = str(e)
                if "429" in error_str:
                    print(f"  ⏳ Rate limit hit, waiting 15 seconds to retry...")
                    time.sleep(15)
                elif "404" in error_str:
                    print(f"  ❌ Model not found, please check SDK config.")
                    break
                else:
                    print(f"  ❌ Error: {e}")
                    break
            
    return all_extracted_data

# --- 4. Local Calculation & Word Export ---
def calculate_and_save(problems: List[ProblemEntry], output_path: str):
    print("🧮 [3/4] Processing answers...")
    doc = Document()
    doc.add_heading('Assessment Answers (NZ Year 4)', 0)
    
    results_map = {}
    for prob in problems:
        if prob.question_type == 'calculation' and prob.expression:
            try:
                # Clean up common non-math chars and execute locally
                exp_to_eval = prob.expression.replace('×','*').replace('÷','/')
                val = eval(exp_to_eval)
                # Format to remove trailing .0
                ans = f"{val:.2f}".rstrip('0').rstrip('.')
                display_text = f"{prob.expression} = {ans}" if prob.is_word_problem else ans
            except:
                display_text = f"{prob.expression} (?)"
        else:
            # If it's multiple choice or true/false, bypass eval() and show the LLM's answer directly
            ans = prob.llm_answer if prob.llm_answer else "(No answer extracted)"
            prefix = prob.question_number + " " if prob.question_number else ""
            display_text = f"{prefix}{ans}"
            
        results_map.setdefault(prob.page_number, []).append(display_text)

    print(f"📝 [4/4] Writing to Word: {output_path}")
    for p_num in sorted(results_map.keys()):
        p = doc.add_paragraph()
        run = p.add_run(f"--- Page {p_num} ---")
        run.bold = True
        
        # Dense layout: separate answers by pipes
        doc.add_paragraph("  |  ".join(results_map[p_num]))
        
    doc.save(output_path)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python batch_solver.py <PDF_PATH>")
        sys.exit(1)
        
    input_pdf = sys.argv[1]
    output_docx = "math_answers_vision.docx"
    
    # Execute pipeline
    img_list = convert_pdf_to_images(input_pdf)
    extracted_data = parse_with_gemini_vision(img_list)

    # ✨ Preview Table
    print("\n" + "="*80)
    print(f"{'Page':<5} | {'Type':<15} | {'Expression/Content':<35} | {'App?'}")
    print("-" * 80)
    for prob in extracted_data:
        word_tag = "✅" if prob.is_word_problem else "❌"
        
        if prob.question_type == 'calculation':
            content = prob.expression
        else:
            content = f"{prob.question_number} {prob.llm_answer}"
            
        # Truncate content if too long for preview
        if len(content) > 35:
            content = content[:32] + "..."
        print(f"{prob.page_number:<5} | {prob.question_type[:14]:<15} | {content:<35} | {word_tag}")
    print("="*80 + "\n")

    # Ask for confirmation before generating docx
    confirm = input("Data mapped above. Generate Word document? (y/n): ")
    if confirm.lower() == 'y':
        calculate_and_save(extracted_data, output_docx)
        print(f"\n✨ Success! Answers saved to: {output_docx}")
    else:
        print("Export cancelled.")
